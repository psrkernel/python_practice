import json
import time
import requests
import docx

path = "F:/code/yequbiancheng/documents/tedTranslate/ted.docx"
paragraphs = docx.Document(path).paragraphs
docxNew = docx.Document()

url = "https://fanyi.youdao.com/translate"
header = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/111.0.0.0 Safari/537.36"}
for paragraph in paragraphs:
    par_text = paragraph.text
    if par_text != "":
        time.sleep(1)
        word = {"i":par_text, "doctype":"json"}
        response = requests.post(url, data = word, headers = header)
        data = json.loads(response.text)
        translate = data["translateResult"][0][0]["tgt"]
        docxNew.add_paragraph(par_text)
        docxNew.add_paragraph(translate)
docxNew.save("F:/code/yequbiancheng/documents/tedTranslate/ted_learn.docx")